from flask import Flask, render_template, request
import docx
import fitz
import pytesseract
from PIL import Image
import re
import os
import docx
import time
import openai
from fpdf import FPDF


app = Flask(__name__)

pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'

def ocr_image(image):
    return pytesseract.image_to_string(image, lang='eng')

def clean_text(text):
    # Remove extra line breaks and multiple spaces
    cleaned_text = re.sub(r'\n+', ' ', text)
    cleaned_text = re.sub(r' +', ' ', cleaned_text)

    # Remove page breaks
    cleaned_text = cleaned_text.replace('\x0c', '')

    # Add a line break before each point (number followed by a full stop)
    cleaned_text = re.sub(r'(\d+\.)', r'\n\1', cleaned_text)

    return cleaned_text.strip()

def process_uploaded_files(files):
    try:
        # Initialize a list to store the Word documents
        doc_list = []

        for file in files:
            # Create a new Word document for each file
            doc = docx.Document()

            # Check if the file is a PDF or an image file (e.g., JPG, PNG)
            if file.filename.lower().endswith('.pdf'):
                # Open the PDF file using PyMuPDF (fitz)
                pdf_document = fitz.open(stream=file.read(), filetype='pdf')

                # Initialize a variable to store the content of all pages as a single paragraph
                full_text = ""

                # Process each page
                for page_num in range(pdf_document.page_count):
                    page = pdf_document[page_num]

                    # Render the page as an image
                    image = page.get_pixmap()
                    img = Image.frombytes("RGB", [image.width, image.height], image.samples)

                    # Perform OCR on the image to extract text
                    text = ocr_image(img)

                    # Clean the extracted text
                    cleaned_text = clean_text(text)

                    # Append the cleaned text to the full_text variable
                    full_text += cleaned_text + " "

                pdf_document.close()
            elif file.filename.lower().endswith(('.jpg', '.jpeg', '.png')):
                # Perform OCR on the image file
                img = Image.open(file)

                # Perform OCR on the image to extract text
                text = ocr_image(img)

                # Clean the extracted text
                cleaned_text = clean_text(text)

                # Append the cleaned text to the full_text variable
                full_text += cleaned_text + " "

            # Add the full text as a single paragraph in the Word document
            doc.add_paragraph(full_text)

            # Append the document to the list
            doc_list.append(doc)

        return doc_list
    except Exception as e:
        print(f"Error while processing uploaded files: {e}")
        return None


# Function to extract text from a docx file
def extract_docx_text(docx_file):
    text = ""
    if os.path.exists(docx_file):
        doc = docx.Document(docx_file)

        for paragraph in doc.paragraphs:
            text += paragraph.text + " "

    return text.strip()

# Path to the reference folder
reference_folder = 'reference'

# Filename of the docx file to process
filename = 'Ref_1.docx'

# Full path to the docx file
docx_file_path = os.path.join(reference_folder, filename)

# Perform text extraction on the specified docx file
extracted_text = extract_docx_text(docx_file_path)

openai.api_key = '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'

def get_gpt_response(prompt, max_tokens=100):
    try:
        response = openai.Completion.create(
            engine="text-davinci-002",  # GPT-3.5 engine
            prompt=prompt,
            max_tokens=max_tokens,  # Adjust the response length as needed
            n=1,  # Number of responses to generate
            stop=None,  # Stop generation at a specific sequence (optional)
            temperature=0.7,  # Controls the randomness of the response
            timeout=30,  # Optional timeout (in seconds)
        )
        return response['choices'][0]['text'].strip()
    except Exception as e:
        print(f"Error: {e}")
        return None

def split_text_into_chunks(text, chunk_size=200):
    words = text.split()
    chunks = [words[i:i + chunk_size] for i in range(0, len(words), chunk_size)]
    return [" ".join(chunk) for chunk in chunks]

def process_long_content(extracted_text):
    response_segments = []

    chunks = split_text_into_chunks(extracted_text, chunk_size=200)
    for chunk in chunks:
        #prompt_text = "Provide the written statement for the given content. I want the supporting evidence in favor of the company and to deny the claims of petitioner too and ask them to provide proof for cross examination at court and all should be done within 150 words for:\n\n"
        prompt_text="Write Written Statement on behalf of an insurance company denying the facts of the given claim by third part all should be done in less than 150 words for:\n\n"
        prompt_text += chunk

        response_text = get_gpt_response(prompt_text)
        if response_text:
            response_segments.append(response_text)
        time.sleep(20)  # Sleep for 20 seconds between iterations

    return "\n".join(response_segments)

# Assuming 'extracted_text' contains the full extracted content
response_text = process_long_content(extracted_text)



# Function to generate and save a PDF document
def save_pdf_to_written_statement_folder(content):
    # Create a PDF document
    pdf = FPDF()
    pdf.add_page()

    # Set font and size for the content
    pdf.set_font("Arial", size=12)

    # Add content to the PDF document
    pdf.multi_cell(0, 10, txt=content, align="L")

    # Path to the "written_statement" folder
    written_statement_folder = 'written_statement'

    # Create the "written_statement" folder if not present
    if not os.path.exists(written_statement_folder):
        os.makedirs(written_statement_folder)

    # Generate the output filename as "WS.pdf"
    output_filename = os.path.join(written_statement_folder, 'WS.pdf')

    # Save the PDF to the specified output filename
    pdf.output(output_filename)

    return output_filename


@app.route('/')
def index():
    return render_template('index.html')

@app.route('/home')
def home():
    return render_template('index.html')

@app.route('/input')
def input_page():
    return render_template('input.html')

@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/uploadfiles', methods=['POST'])
def uploadfiles():
    casetype = request.form['caseType']
    uploaded_files = request.files.getlist('files[]')

    # Process uploaded files and get the list of Word documents
    result_docs = process_uploaded_files(uploaded_files)

    if result_docs:
        # Create a folder named "reference" if it doesn't exist
        reference_folder = 'reference'
        if not os.path.exists(reference_folder):
            os.makedirs(reference_folder)

        # Save each Word document to the "reference" folder
        for i, doc in enumerate(result_docs):
            # Generate a unique filename for the Word document
            filename = f"Ref_{i+1}.docx"
            output_path = os.path.join(reference_folder, filename)

            # Save the Word document to the specified output path
            doc.save(output_path)
            print(f"Scanned PDFs/Images converted to Word and saved to: {output_path}")

        # Assuming 'response_text' contains the full extracted content
        extracted_text = "\n\n".join(extract_docx_text(os.path.join(reference_folder, f"Ref_{i+1}.docx")) for i in range(len(result_docs)))

        # Process the extracted text using GPT-3
        response_text = process_long_content(extracted_text)

        # Save the response to a PDF
        pdf_filename = save_pdf_to_written_statement_folder(response_text)

        print(f"PDF saved to: {pdf_filename}")

        return render_template('upload.html', message="Processing completed. PDF saved!")
    else:
        return render_template('upload.html', message="Error processing files. Please try again.")

if __name__ == '__main__':
    app.run(debug=True)
